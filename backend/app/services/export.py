"""
匯出服務
"""
from io import BytesIO
from datetime import datetime

from app.models.transcript import Transcript


def export_to_markdown(transcript: Transcript) -> str:
    """匯出為 Markdown 格式"""
    lines = []

    # 標題
    lines.append(f"# {transcript.title or '轉錄記錄'}")
    lines.append("")
    lines.append(f"**建立時間**: {transcript.created_at.strftime('%Y-%m-%d %H:%M')}")
    if transcript.language:
        lines.append(f"**語言**: {transcript.language}")
    if transcript.audio_duration:
        mins = int(transcript.audio_duration // 60)
        secs = int(transcript.audio_duration % 60)
        lines.append(f"**時長**: {mins}:{secs:02d}")
    lines.append("")

    # 摘要
    if transcript.summary:
        lines.append("---")
        lines.append("")
        lines.append(transcript.summary)
        lines.append("")

    # 轉錄內容
    lines.append("---")
    lines.append("")
    lines.append("## 完整轉錄")
    lines.append("")

    if transcript.transcript_segments:
        for seg in transcript.transcript_segments:
            start = seg.get("start", 0)
            end = seg.get("end", 0)
            text = seg.get("text", "")
            start_str = f"{int(start // 60):02d}:{int(start % 60):02d}"
            end_str = f"{int(end // 60):02d}:{int(end % 60):02d}"
            lines.append(f"**[{start_str} - {end_str}]** {text}")
            lines.append("")
    else:
        lines.append(transcript.transcript_text or "")

    return "\n".join(lines)


def export_to_pdf(transcript: Transcript) -> bytes:
    """匯出為 PDF 格式"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    buffer = BytesIO()

    # 設定 PDF
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2*cm,
        leftMargin=2*cm,
        topMargin=2*cm,
        bottomMargin=2*cm,
    )

    # 嘗試註冊中文字體
    try:
        # macOS 系統字體
        pdfmetrics.registerFont(TTFont("PingFang", "/System/Library/Fonts/PingFang.ttc"))
        font_name = "PingFang"
    except Exception:
        font_name = "Helvetica"

    # 樣式
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "Title",
        parent=styles["Heading1"],
        fontName=font_name,
        fontSize=18,
        spaceAfter=20,
    )
    normal_style = ParagraphStyle(
        "Normal",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=11,
        leading=16,
    )
    heading_style = ParagraphStyle(
        "Heading",
        parent=styles["Heading2"],
        fontName=font_name,
        fontSize=14,
        spaceBefore=15,
        spaceAfter=10,
    )

    # 內容
    story = []

    # 標題
    story.append(Paragraph(transcript.title or "轉錄記錄", title_style))

    # 資訊
    info_text = f"建立時間: {transcript.created_at.strftime('%Y-%m-%d %H:%M')}"
    if transcript.language:
        info_text += f" | 語言: {transcript.language}"
    if transcript.audio_duration:
        mins = int(transcript.audio_duration // 60)
        secs = int(transcript.audio_duration % 60)
        info_text += f" | 時長: {mins}:{secs:02d}"
    story.append(Paragraph(info_text, normal_style))
    story.append(Spacer(1, 20))

    # 摘要
    if transcript.summary:
        story.append(Paragraph("摘要", heading_style))
        # 處理摘要中的換行
        for line in transcript.summary.split("\n"):
            if line.strip():
                story.append(Paragraph(line, normal_style))
        story.append(Spacer(1, 15))

    # 轉錄內容
    story.append(Paragraph("完整轉錄", heading_style))
    if transcript.transcript_text:
        # 分段顯示
        paragraphs = transcript.transcript_text.split("\n")
        for para in paragraphs:
            if para.strip():
                story.append(Paragraph(para, normal_style))
                story.append(Spacer(1, 6))

    doc.build(story)
    return buffer.getvalue()


def export_to_docx(transcript: Transcript) -> bytes:
    """匯出為 Word 格式"""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 標題
    title = doc.add_heading(transcript.title or "轉錄記錄", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 資訊
    info_para = doc.add_paragraph()
    info_para.add_run(f"建立時間: {transcript.created_at.strftime('%Y-%m-%d %H:%M')}")
    if transcript.language:
        info_para.add_run(f" | 語言: {transcript.language}")
    if transcript.audio_duration:
        mins = int(transcript.audio_duration // 60)
        secs = int(transcript.audio_duration % 60)
        info_para.add_run(f" | 時長: {mins}:{secs:02d}")

    doc.add_paragraph()

    # 摘要
    if transcript.summary:
        doc.add_heading("摘要", level=1)
        for line in transcript.summary.split("\n"):
            if line.strip():
                doc.add_paragraph(line)
        doc.add_paragraph()

    # 轉錄內容
    doc.add_heading("完整轉錄", level=1)

    if transcript.transcript_segments:
        for seg in transcript.transcript_segments:
            start = seg.get("start", 0)
            end = seg.get("end", 0)
            text = seg.get("text", "")
            start_str = f"{int(start // 60):02d}:{int(start % 60):02d}"
            end_str = f"{int(end // 60):02d}:{int(end % 60):02d}"

            para = doc.add_paragraph()
            run = para.add_run(f"[{start_str} - {end_str}] ")
            run.bold = True
            para.add_run(text)
    else:
        if transcript.transcript_text:
            for para_text in transcript.transcript_text.split("\n"):
                if para_text.strip():
                    doc.add_paragraph(para_text)

    # 儲存到 buffer
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
